import sys
import openpyxl
from docx import Document

if len(sys.argv) < 5:
    print("Usage: python script.py excel_file doc_file start_row count_rows")
    sys.exit(1)
excel_file = sys.argv[1]
doc_file = sys.argv[2]
start_row = int(sys.argv[3])
count_rows = int(sys.argv[4])

wb = openpyxl.load_workbook(excel_file)
ws = wb.active

doc = Document(doc_file)
doc_name = doc_file.split(".")[0]

header_row = ws[1]
tmp_word_positions = [cell for i, cell in enumerate(header_row) if cell.value]

for row in ws.iter_rows(min_row=start_row, max_row=start_row+count_rows-1, values_only=True):
    if any(row):
        new_doc = Document(doc_file)
        replaceCounter = 0
        for i, pos in enumerate(tmp_word_positions):
            tmp_word = header_row[i].value
            new_value = str(row[i])
            #print(f"Replacing '{tmp_word}' with '{new_value}'")
            for paragraph in new_doc.paragraphs:
                if tmp_word in paragraph.text:
                    #print(f"  Found '{tmp_word}' in paragraph '{paragraph.text}'")
                    inline = paragraph.runs
                    for il in inline:
                        if tmp_word in il.text:
                            #print(f"    Found '{tmp_word}' in inline run '{il.text}'")
                            il.text = il.text.replace(tmp_word, new_value)
                            replaceCounter+=1

            for table in new_doc.tables:
                for r in table.rows:
                    for cell in r.cells:
                        for paragraph in cell.paragraphs:
                            if tmp_word in paragraph.text:
                                #print(f"  Found '{tmp_word}' in paragraph '{paragraph.text}'")
                                for inline in paragraph.runs:
                                    if tmp_word in inline.text:
                                        inline.text = inline.text.replace(tmp_word, new_value)
                                        replaceCounter+=1

        filename = doc_name+" "+row[1]
        print("Replaced {} fields in {}".format(replaceCounter, filename))
        new_doc.save(f'{filename}.docx')